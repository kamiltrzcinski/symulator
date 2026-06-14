from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "przegladarka-pojazdow.docx"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="Lista punktowana")
    paragraph.add_run(text)
    return paragraph


def add_step(document, text):
    paragraph = document.add_paragraph(style="Lista numerowana")
    paragraph.add_run(text)
    return paragraph


document = Document()
section = document.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = document.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name, base_name in (
    ("Lista punktowana", "List Bullet"),
    ("Lista numerowana", "List Number"),
):
    style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base_name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

header = section.header.paragraphs[0]
header.text = "SYMULATOR | NARZĘDZIA DEWELOPERSKIE"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.name = "Calibri"
header.runs[0].font.size = Pt(8.5)
header.runs[0].font.color.rgb = MUTED

footer = section.footer.paragraphs[0]
footer.text = "Edytor taboru i składów | instrukcja użytkownika"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.name = "Calibri"
footer.runs[0].font.size = Pt(8.5)
footer.runs[0].font.color.rgb = MUTED

title = document.add_paragraph()
title.paragraph_format.space_before = Pt(42)
title.paragraph_format.space_after = Pt(6)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Edytor taboru i składów")
run.font.name = "Calibri"
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = DARK_BLUE

subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(24)
run = subtitle.add_run("Podręcznik użytkownika narzędzia vehicle-browser")
run.font.name = "Calibri"
run.font.size = Pt(14)
run.font.color.rgb = MUTED

lead = document.add_table(rows=1, cols=1)
lead.autofit = False
lead.columns[0].width = Inches(6.5)
cell = lead.cell(0, 0)
cell.width = Inches(6.5)
set_cell_shading(cell, "E8EEF5")
set_cell_margins(cell, 140, 180, 140, 180)
paragraph = cell.paragraphs[0]
paragraph.paragraph_format.space_after = Pt(0)
run = paragraph.add_run(
    "Narzędzie służy do przeglądania i edycji danych taborowych, tworzenia pojazdów "
    "oraz otwierania i układania składów z packages/ lub symulator-data."
)
run.font.name = "Calibri"
run.font.size = Pt(11)
run.font.bold = True

document.add_heading("1. Przeznaczenie", level=1)
document.add_paragraph(
    "Edytor taboru i składów jest aplikacją Qt6 dla twórców i opiekunów danych. "
    "Nie jest częścią interfejsu gracza. Łączy katalog typów pojazdów, listę "
    "konkretnych egzemplarzy oraz prosty kreator składów."
)
for item in (
    "przeglądanie, filtrowanie i sortowanie typów pojazdów;",
    "wyświetlanie pojazdów przypisanych do wybranego typu;",
    "tworzenie i edycję plików vehicle.json bez zmiany istniejącego UID;",
    "otwieranie składów i zmianę kolejności metodą przeciągnij i upuść;",
    "nadpisywanie składów lub zapisywanie kopii z nowym UID;",
    "podgląd struktury UID w polskiej legendzie.",
):
    add_bullet(document, item)

document.add_heading("2. Uruchomienie i źródła danych", level=1)
document.add_paragraph(
    "Aplikację należy uruchamiać z katalogu głównego repozytorium symulator. "
    "Domyślnie czyta dane z katalogu packages/."
)
code = document.add_paragraph()
code.paragraph_format.left_indent = Inches(0.25)
code.paragraph_format.space_after = Pt(8)
run = code.add_run("vehicle-browser [--data-dir <katalog>] [--help]")
run.font.name = "Consolas"
run.font.size = Pt(10)

table = document.add_table(rows=1, cols=2)
table.autofit = False
table.columns[0].width = Inches(1.8)
table.columns[1].width = Inches(4.7)
headers = ("Tryb", "Zastosowanie")
for index, text in enumerate(headers):
    table.cell(0, index).text = text
    set_cell_shading(table.cell(0, index), "E8EEF5")
for mode, description in (
    ("packages/", "Dane wydane przez symulator-data i pobrane do repozytorium symulator."),
    ("--data-dir", "Bezpośrednia praca na dowolnym katalogu JSON, np. symulator-data\\data."),
    ("Plik > Otwórz katalog...", "Zmiana źródła bez ponownego uruchamiania programu."),
):
    cells = table.add_row().cells
    cells[0].text = mode
    cells[1].text = description
for row in table.rows:
    for cell in row.cells:
        set_cell_margins(cell)
        cell.vertical_alignment = 1

document.add_heading("3. Przeglądanie danych", level=1)
add_step(document, "W lewym panelu wybierz typ pojazdu.")
add_step(document, "Użyj pola Filtruj typy pojazdów..., aby zawęzić listę.")
add_step(document, "Kliknij nagłówek kolumny, aby zmienić sortowanie.")
add_step(document, "Sprawdź pojazdy przypisane do typu w środkowym panelu.")
document.add_paragraph(
    "Po wybraniu typu środkowa tabela pokazuje tylko pojazdy o pasującym type_uid. "
    "Usunięcie zaznaczenia przywraca pełną listę."
)

document.add_heading("4. Tworzenie pojazdu", level=1)
for text in (
    "Wybierz typ pojazdu.",
    "Kliknij Nowy pojazd....",
    "Wpisz numer boczny. Opcjonalnie dodaj UID przewoźnika, numer inwentarzowy i uwagi.",
    "Kliknij Zapisz i wskaż katalog docelowy.",
):
    add_step(document, text)
document.add_paragraph(
    "Program proponuje UID rodzaju VEHICLE. Przed zapisem ponownie sprawdza "
    "jego dostępność i w razie kolizji wybiera kolejną wolną instancję."
)
code = document.add_paragraph()
run = code.add_run("<katalog>/<numer-boczny>/vehicle.json")
run.font.name = "Consolas"
run.font.size = Pt(10)

document.add_heading("5. Edycja istniejącego pojazdu", level=1)
for text in (
    "Zaznacz pojazd w środkowym panelu.",
    "Kliknij Edytuj pojazd... albo dwukrotnie kliknij wiersz.",
    "Zmień dane formularza i kliknij Zapisz.",
):
    add_step(document, text)
document.add_paragraph(
    "Edytor nadpisuje plik źródłowy, zachowuje UID i nie usuwa dodatkowych pól "
    "JSON, których formularz nie wyświetla. Dane są odświeżane automatycznie."
)

document.add_heading("6. Tworzenie i edycja składu", level=1)
for text in (
    "Wybierz istniejący skład z listy albo pozycję (Nowy skład).",
    "Zaznacz pojazd i kliknij Dodaj do składu.",
    "Dodaj pozostałe pojazdy.",
    "Przeciągaj pozycje, aby ustawić ich kolejność; użyj Usuń pojazd dla błędnej pozycji.",
    "Uzupełnij numer, nazwę, kategorię i opcjonalny UID przewoźnika.",
    "Kliknij Zapisz, aby nadpisać plik, albo Zapisz jako..., aby utworzyć nowy UID.",
):
    add_step(document, text)
document.add_paragraph(
    "Nowy skład otrzymuje bezkolizyjny UID rodzaju TRAIN_CONSIST. Edytowany skład "
    "zachowuje swój UID i dodatkowe pola JSON. Lista vehicle_uids jest zapisywana "
    "w kolejności widocznej w prawym panelu."
)

document.add_heading("7. Współpraca z symulator-data", level=1)
document.add_paragraph(
    "Kod aplikacji znajduje się w repozytorium symulator, natomiast dane "
    "taborowe należą logicznie do symulator-data. Edytor czyta oba "
    "warianty danych i zapisuje pliki zgodne z ich schematami."
)
for item in (
    "Program nie wykonuje commitów ani operacji Git.",
    "Nowe pliki należy zapisać w odpowiednim katalogu symulator-data.",
    "Po zapisie dane są odświeżane automatycznie; warto zweryfikować rekord w tabeli.",
    "Zmiany należy zatwierdzić standardowym procesem przeglądu danych.",
):
    add_bullet(document, item)

document.add_heading("8. Legenda UID", level=1)
document.add_paragraph(
    "Polecenie Pomoc > Legenda UID pokazuje układ bitów, domenę, rodzaj oraz "
    "znaczenie pola ZAKRES. Legenda działa nawet wtedy, gdy źródło danych nie "
    "zostało wczytane."
)

document.add_heading("9. Rozwiązywanie problemów", level=1)
issues = (
    ("Brak pluginu windows", "Umieść qwindows.dll w katalogu platforms obok programu; dla Debug użyj qwindowsd.dll."),
    ("Brak danych", "Uruchom program z katalogu symulator albo użyj --data-dir."),
    ("Brak pojazdu w składzie", "Wczytaj kompletne źródło danych albo popraw vehicle_uids."),
    ("Nie można zapisać składu", "Dodaj co najmniej jeden pojazd i sprawdź uprawnienia do katalogu docelowego."),
)
table = document.add_table(rows=1, cols=2)
table.autofit = False
table.columns[0].width = Inches(2.0)
table.columns[1].width = Inches(4.5)
table.cell(0, 0).text = "Problem"
table.cell(0, 1).text = "Rozwiązanie"
for cell in table.rows[0].cells:
    set_cell_shading(cell, "E8EEF5")
for problem, solution in issues:
    cells = table.add_row().cells
    cells[0].text = problem
    cells[1].text = solution
for row in table.rows:
    for cell in row.cells:
        set_cell_margins(cell)
        cell.vertical_alignment = 1

document.add_paragraph()
closing = document.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run("Dokumentacja dla wersji narzędzia 1.0")
run.font.name = "Calibri"
run.font.size = Pt(9)
run.font.color.rgb = MUTED

document.save(OUTPUT)
print(OUTPUT)
